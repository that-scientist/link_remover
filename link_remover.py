#!/usr/bin/env python3
"""
Link Remover - Removes hyperlinks from .docx and PDF files
while preserving human-readable text.
"""

import sys
import os
import shutil
import subprocess
from pathlib import Path

# Auto-detect and use venv Python if available
_script_dir = Path(__file__).parent.resolve()
_venv_dir = _script_dir / 'venv'
_venv_python = _venv_dir / 'bin' / 'python3'
_requirements_file = _script_dir / 'requirements.txt'

def ensure_venv():
    """Create virtual environment if it doesn't exist."""
    if not _venv_dir.exists():
        print("Creating virtual environment...")
        try:
            subprocess.run([sys.executable, '-m', 'venv', str(_venv_dir)], 
                         check=True, capture_output=True)
            print("✓ Virtual environment created")
            return True
        except subprocess.CalledProcessError as e:
            print(f"Error creating virtual environment: {e}")
            return False
    return True

def install_requirements():
    """Install packages from requirements.txt."""
    if not _requirements_file.exists():
        print(f"Warning: {_requirements_file} not found")
        return False
    
    print("Installing required packages...")
    pip_cmd = [sys.executable, '-m', 'pip', 'install', '-q', '-r', str(_requirements_file)]
    
    try:
        subprocess.run(pip_cmd, check=True, capture_output=True, text=True)
        print("✓ Packages installed successfully")
        return True
    except subprocess.CalledProcessError as e:
        error_msg = e.stderr if e.stderr else (e.stdout if e.stdout else 'Unknown error')
        print(f"Error installing packages: {error_msg}")
        print(f"Try running manually: {sys.executable} -m pip install -r {_requirements_file}")
        return False

def check_and_install_packages():
    """Check if required packages are installed, install if missing."""
    # Check if packages are installed
    docx_available = False
    pdf_available = False
    
    try:
        import docx  # noqa: F401
        docx_available = True
    except ImportError:
        pass
    
    try:
        import pypdf  # noqa: F401
        pdf_available = True
    except ImportError:
        try:
            import PyPDF2  # noqa: F401
            pdf_available = True
        except ImportError:
            pass
    
    # If both are available, we're good
    if docx_available and pdf_available:
        return True
    
    # Some packages are missing, try to install
    if install_requirements():
        # Try imports again after installation
        try:
            import docx  # noqa: F401
            docx_available = True
        except ImportError:
            pass
        
        try:
            import pypdf  # noqa: F401
            pdf_available = True
        except ImportError:
            try:
                import PyPDF2  # noqa: F401
                pdf_available = True
            except ImportError:
                pass
        
        return docx_available and pdf_available
    
    return False

# Setup virtual environment and ensure we're using it
if not sys.executable.startswith(str(_script_dir / 'venv')):
    # Not using venv, check if we should create and use it
    if ensure_venv() and _venv_python.exists():
        # Re-execute with venv Python
        os.execv(str(_venv_python), [str(_venv_python)] + sys.argv)
else:
    # We're in venv, check and install packages if needed
    # This will install packages if missing, but won't re-execute
    # Packages will be available on next import attempt
    check_and_install_packages()

# Now try to import required packages
try:
    from docx import Document
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader, PdfWriter
    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader, PdfWriter
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False


def remove_hyperlinks_from_docx(input_path: Path, output_path: Path) -> bool:
    """
    Remove hyperlinks from a .docx file while preserving text.
    
    Args:
        input_path: Path to input .docx file
        output_path: Path to save output .docx file
        
    Returns:
        True if successful, False otherwise
    """
    if not DOCX_AVAILABLE:
        # Try to install packages if we're in venv, then re-execute
        if sys.executable.startswith(str(_script_dir / 'venv')):
            print("python-docx library not found. Attempting to install...")
            if install_requirements():
                # Re-execute script to pick up newly installed packages
                print("✓ Packages installed. Restarting script...")
                os.execv(sys.executable, [sys.executable] + sys.argv)
            else:
                print("Error: Failed to install required packages.")
                return False
        else:
            print("Error: python-docx library not installed.")
            print(f"  Current Python: {sys.executable}")
            print(f"  Please run the script again to auto-install packages, or install manually:")
            print(f"    {sys.executable} -m pip install python-docx")
            return False
    
    try:
        doc = Document(input_path)
        
        def process_paragraph(paragraph):
            """Process a paragraph to remove hyperlinks while preserving text and setting color to black."""
            # Find all hyperlinks in the paragraph using findall with qualified name
            hyperlink_qn = qn('w:hyperlink')
            run_qn = qn('w:r')
            color_qn = qn('w:color')
            rPr_qn = qn('w:rPr')
            
            # Use iter to find all hyperlinks recursively
            for element in paragraph._element.iter():
                if element.tag == hyperlink_qn:
                    hyperlink = element
                    # Get the parent element
                    parent = hyperlink.getparent()
                    if parent is not None:
                        # Get all child runs from the hyperlink
                        runs = list(hyperlink.findall(run_qn))
                        
                        # Process each run to set color to black
                        for run_elem in runs:
                            # Get or create run properties
                            rPr = run_elem.find(rPr_qn)
                            if rPr is None:
                                rPr = run_elem.makeelement(rPr_qn)
                                run_elem.insert(0, rPr)
                            
                            # Remove existing color element if it exists
                            color_elem = rPr.find(color_qn)
                            if color_elem is not None:
                                rPr.remove(color_elem)
                            
                            # Set color to black (000000 in hex, which is "auto" or black)
                            # In Word, "auto" (000000) means use the default text color (black)
                            color_elem = run_elem.makeelement(color_qn)
                            color_elem.set(qn('w:val'), '000000')  # Black color
                            rPr.append(color_elem)
                        
                        # Insert runs before the hyperlink, then remove it
                        hyperlink_index = list(parent).index(hyperlink)
                        
                        # Move runs out of hyperlink
                        for run_elem in reversed(runs):
                            parent.insert(hyperlink_index, run_elem)
                        
                        # Remove the hyperlink element
                        parent.remove(hyperlink)
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph)
        
        # Process tables for hyperlinks
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error processing {input_path}: {e}")
        import traceback
        traceback.print_exc()
        return False


def remove_hyperlinks_from_pdf(input_path: Path, output_path: Path) -> bool:
    """
    Remove hyperlinks from a PDF file while preserving text.
    
    Args:
        input_path: Path to input PDF file
        output_path: Path to save output PDF file
        
    Returns:
        True if successful, False otherwise
    """
    if not PDF_AVAILABLE:
        # Try to install packages if we're in venv, then re-execute
        if sys.executable.startswith(str(_script_dir / 'venv')):
            print("pypdf library not found. Attempting to install...")
            if install_requirements():
                # Re-execute script to pick up newly installed packages
                print("✓ Packages installed. Restarting script...")
                os.execv(sys.executable, [sys.executable] + sys.argv)
            else:
                print("Error: Failed to install required packages.")
                return False
        else:
            print("Error: pypdf library not installed.")
            print(f"  Current Python: {sys.executable}")
            print(f"  Please run the script again to auto-install packages, or install manually:")
            print(f"    {sys.executable} -m pip install pypdf")
            return False
    
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            # Clone the page
            page_obj = writer.add_page(page)
            
            # Remove annotations (which include hyperlinks) from the page dictionary
            if '/Annots' in page_obj:
                # Access the underlying page dictionary
                page_dict = page_obj.get_object()
                if '/Annots' in page_dict:
                    del page_dict['/Annots']
        
        # Write the output PDF
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        return True
        
    except Exception as e:
        print(f"Error processing {input_path}: {e}")
        import traceback
        traceback.print_exc()
        return False


def process_file(input_file: Path, input_dir: Path, output_dir: Path, done_dir: Path) -> bool:
    """
    Process a single file: remove hyperlinks and move to done folder.
    
    Args:
        input_file: Path to the file to process
        input_dir: Input directory path
        output_dir: Output directory path
        done_dir: Done directory path
        
    Returns:
        True if successful, False otherwise
    """
    file_ext = input_file.suffix.lower()
    file_name = input_file.name
    
    # Determine output filename
    output_file = output_dir / file_name
    
    # Process based on file type
    success = False
    if file_ext == '.docx':
        print(f"Processing .docx file: {file_name}")
        success = remove_hyperlinks_from_docx(input_file, output_file)
    elif file_ext == '.pdf':
        print(f"Processing PDF file: {file_name}")
        success = remove_hyperlinks_from_pdf(input_file, output_file)
    else:
        print(f"Skipping unsupported file type: {file_name}")
        return False
    
    if success:
        # Move original file to done folder
        done_file = done_dir / file_name
        shutil.move(str(input_file), str(done_file))
        print(f"✓ Processed and moved {file_name} to done folder")
        return True
    else:
        print(f"✗ Failed to process {file_name}")
        return False


def main():
    """Main function to process all files in the input folder."""
    # Get the script directory
    script_dir = Path(__file__).parent
    
    # Define directories
    input_dir = script_dir / 'input'
    output_dir = script_dir / 'out'
    done_dir = script_dir / 'done'
    
    # Create directories if they don't exist
    input_dir.mkdir(exist_ok=True)
    output_dir.mkdir(exist_ok=True)
    done_dir.mkdir(exist_ok=True)
    
    # Check for files in input directory
    input_files = list(input_dir.glob('*.docx')) + list(input_dir.glob('*.pdf'))
    
    if not input_files:
        print("No .docx or .pdf files found in the input folder.")
        return
    
    print(f"Found {len(input_files)} file(s) to process.\n")
    
    # Process each file
    success_count = 0
    for input_file in input_files:
        if process_file(input_file, input_dir, output_dir, done_dir):
            success_count += 1
        print()
    
    print(f"Processing complete: {success_count}/{len(input_files)} files processed successfully.")


if __name__ == '__main__':
    main()

